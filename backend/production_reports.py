from datetime import datetime, timedelta
from io import BytesIO
import re

from production import build_operator_summary, build_production_kpi_payload, filter_production_period


def safe_number(value, digits=1):
    return None if value is None else round(float(value), digits)


def build_production_anomalies(payload):
    anomalies = []
    for item in payload.get("rows", []):
        checks = [
            ("time_booked_today", item.get("time_booked_today"), lambda value: value < 0 or value > 200, "Expected between 0 and 200 hours", "Check the total booked-time calculation and decimal placement."),
            ("utilisation_percentage", item.get("utilisation_percentage"), lambda value: value < 0 or value > 100, "Expected between 0% and 100%", "Check the plant utilisation calculation and source bookings."),
            ("labour_hours", item.get("labour_hours"), lambda value: value < 0 or value > 200, "Expected between 0 and 200 hours", "Check labour-hour entries for the day."),
        ]
        for field, value, invalid, reason, action in checks:
            if value is not None and invalid(float(value)):
                anomalies.append({"type": "Production", "severity": "Review", "date": item.get("date"), "name": "Plant", "record_id": item.get("record_id"), "field": field, "value": value, "reason": reason, "action": action})
        for label, key in (("Sort", "sort_throughput"), ("Grind", "grind_throughput"), ("Press", "press_throughput"), ("Trim", "trim_throughput")):
            if item.get(key) is None:
                anomalies.append({"type": "Production", "severity": "Information", "date": item.get("date"), "name": label, "record_id": item.get("record_id"), "field": key, "value": None, "reason": "No throughput value recorded", "action": "Confirm this department had no activity or correct the missing value."})

    for item in payload.get("operator_rows", []):
        checks = [
            ("booked_hours", item.get("booked_hours"), lambda value: value < 0 or value > 24, "Booked time is negative or above 24 hours", "Check the operator booking entries and decimal placement."),
            ("clocked_hours", item.get("clocked_hours"), lambda value: value < 0 or value > 24, "Clocked time is negative or above 24 hours", "Check the time-clock entry."),
            ("productivity", item.get("productivity"), lambda value: value < 0 or value > 250, "Productivity is outside 0–250%", "Confirm whether unusually high productivity is valid or caused by booking data."),
            ("target_achieved", item.get("target_achieved"), lambda value: value < 0 or value > 250, "Target achieved is outside 0–250%", "Check the LF target calculation and applicability for this operator."),
        ]
        for field, value, invalid, reason, action in checks:
            if value is not None and invalid(float(value)):
                anomalies.append({"type": "Operator", "severity": "Review", "date": item.get("date"), "name": item.get("name"), "record_id": item.get("record_id"), "field": field, "value": value, "reason": reason, "action": action})
        if (item.get("booked_hours") or 0) > 0 and not item.get("clocked_hours"):
            anomalies.append({"type": "Operator", "severity": "Review", "date": item.get("date"), "name": item.get("name"), "record_id": item.get("record_id"), "field": "clocked_hours", "value": item.get("clocked_hours"), "reason": "Booked time exists without clocked time", "action": "Check the operator clock-in record."})
    return sorted(anomalies, key=lambda item: (str(item.get("date") or ""), item.get("type") or "", item.get("name") or ""), reverse=True)


def build_anomaly_workbook(payload):
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.worksheet.table import Table, TableStyleInfo

    anomalies = build_production_anomalies(payload)
    workbook = Workbook()
    summary = workbook.active
    summary.title = "Summary"
    summary.append(["Production Data Quality Report"])
    summary.append(["Period", period_label(payload)])
    summary.append(["Generated", datetime.now().strftime("%Y-%m-%d %H:%M")])
    summary.append(["Total anomalies", len(anomalies)])
    summary.append(["Review", sum(item["severity"] == "Review" for item in anomalies)])
    summary.append(["Information", sum(item["severity"] == "Information" for item in anomalies)])
    summary.append([])
    summary.append(["Purpose", "Use this report to investigate unusual or missing FileMaker production entries. Values are flagged for review; they are not automatically assumed to be incorrect."])
    summary["A1"].font = Font(size=18, bold=True, color="245CFF")
    summary.column_dimensions["A"].width = 22
    summary.column_dimensions["B"].width = 95
    summary["B8"].alignment = Alignment(wrap_text=True, vertical="top")

    sheet = workbook.create_sheet("Anomalies")
    headers = ["Severity", "Type", "Date", "Person / Area", "Record ID", "Field", "Value", "Why flagged", "Suggested check", "Correction status", "Correction notes"]
    sheet.append(headers)
    for item in anomalies:
        sheet.append([item.get("severity"), item.get("type"), item.get("date"), item.get("name"), item.get("record_id"), item.get("field"), item.get("value"), item.get("reason"), item.get("action"), "Open", ""])
    if anomalies:
        table = Table(displayName="ProductionAnomalies", ref=f"A1:K{len(anomalies) + 1}")
        table.tableStyleInfo = TableStyleInfo(name="TableStyleMedium2", showRowStripes=True, showFirstColumn=False, showLastColumn=False)
        sheet.add_table(table)
    header_fill = PatternFill("solid", fgColor="245CFF")
    for cell in sheet[1]:
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = header_fill
        cell.alignment = Alignment(wrap_text=True, vertical="center")
    widths = [13, 13, 13, 22, 13, 24, 14, 38, 48, 18, 42]
    for index, width in enumerate(widths, start=1):
        sheet.column_dimensions[chr(64 + index)].width = width
    for row in sheet.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    sheet.freeze_panes = "A2"

    definitions = workbook.create_sheet("Rules")
    definitions.append(["Rule", "Reason"])
    rules = [
        ("Booked time below 0 or above 24 hours", "Likely correction, booking or decimal-placement issue."),
        ("Clocked time below 0 or above 24 hours", "Outside a plausible daily clocking range."),
        ("Productivity or target outside 0–250%", "May be valid exceptional performance, but warrants review."),
        ("Plant time above 200 hours", "Potential aggregate calculation or decimal-placement issue."),
        ("Missing department throughput", "Could mean no activity or a missing entry; confirmation is required."),
    ]
    for row in rules:
        definitions.append(row)
    definitions.column_dimensions["A"].width = 45
    definitions.column_dimensions["B"].width = 85
    for cell in definitions[1]:
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = header_fill
    for row in definitions.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    output = BytesIO()
    workbook.save(output)
    return output.getvalue(), len(anomalies)


def summarize_rows(rows, operator_rows):
    revenue = [item.get("production_revenue_today") for item in rows if item.get("production_revenue_today") is not None]
    press_lf = [item.get("press_throughput") for item in rows if item.get("press_throughput") is not None]
    booked = sum(item.get("booked_hours") or 0 for item in operator_rows if (item.get("clocked_hours") or 0) > 0)
    clocked = sum(item.get("clocked_hours") or 0 for item in operator_rows if (item.get("clocked_hours") or 0) > 0)
    return {
        "production_days": len(rows),
        "total_revenue": sum(revenue),
        "average_daily_revenue": sum(revenue) / len(revenue) if revenue else None,
        "total_press_lf": sum(press_lf),
        "average_press_lf": sum(press_lf) / len(press_lf) if press_lf else None,
        "total_mats_repaired_through_presses": sum(sum(item.get(f"ff{press}_mats") or 0 for press in (1, 2, 3)) for item in rows),
        "total_cycles": sum(sum(item.get(f"ff{press}_cycles") or 0 for press in (1, 2, 3)) for item in rows),
        "recook_lf": sum(item.get("recook_lf") or 0 for item in rows),
        "labour_hours": sum(item.get("labour_hours") or 0 for item in rows),
        "plant_productivity": booked / clocked * 100 if clocked else None,
        "latest_backlog_weeks": rows[-1].get("backlog_weeks") if rows else None,
        "department_average_lf": {
            key: safe_number(sum(values) / len(values), 0) if values else None
            for key, values in {
                label: [item.get(field) for item in rows if item.get(field) is not None]
                for label, field in (("Sort", "sort_throughput"), ("Grind", "grind_throughput"), ("Press", "press_throughput"), ("Trim", "trim_throughput"))
            }.items()
        },
    }


def build_ai_report_facts(result, days):
    current_rows, current_operators = filter_production_period(result, days=days)
    if current_rows:
        current_end = datetime.strptime(current_rows[-1]["date"], "%Y-%m-%d")
        current_start = current_end - timedelta(days=max(1, days) - 1)
        previous_end = current_start - timedelta(days=1)
        previous_start = previous_end - timedelta(days=max(1, days) - 1)
        previous_rows = [item for item in result.get("production_rows", []) if previous_start.strftime("%Y-%m-%d") <= item["date"] <= previous_end.strftime("%Y-%m-%d")]
        previous_operators = [item for item in result.get("operator_rows", []) if previous_start.strftime("%Y-%m-%d") <= item["date"] <= previous_end.strftime("%Y-%m-%d")]
    else:
        current_start = current_end = previous_start = previous_end = None
        previous_rows, previous_operators = [], []
    operators = build_operator_summary(current_operators)
    facts = {
        "period": {"start": current_start.strftime("%Y-%m-%d") if current_start else "", "end": current_end.strftime("%Y-%m-%d") if current_end else "", "rolling_days": days},
        "current": summarize_rows(current_rows, current_operators),
        "previous_period": {"start": previous_start.strftime("%Y-%m-%d") if previous_start else "", "end": previous_end.strftime("%Y-%m-%d") if previous_end else "", **summarize_rows(previous_rows, previous_operators)},
        "operators": [{key: safe_number(value) if isinstance(value, float) else value for key, value in item.items()} for item in operators],
        "business_context": {
            "preferred_backlog_weeks": 2.0,
            "backlog_interpretation": "Backlog below approximately two weeks means production has less new-order repair work available and should prompt sales activity to secure more orders. Growth toward two weeks is positive, not a turnaround concern. A low backlog normally enables faster-than-usual customer turnaround.",
        },
        "data_notes": ["Today is excluded.", "Backlog contains new customer orders awaiting their first repair only; re-cook work is never included in backlog.", "The preferred backlog is approximately two weeks. A backlog below that level can constrain production because only received customer mats can be repaired, and should prompt the sales team to secure more orders.", "Backlog growth that remains below approximately two weeks is movement toward a healthier workload and must not be described as harming turnaround. Low backlog normally means orders can be returned faster than usual.", "Re-cook activity is a separate quality measure and is separate from press throughput.", "Press throughput is FF1 + FF2 + FF3 total LF.", "Trudy Dunlap, Kelly Bainbridge, Lois Horace, Temp1 and Temp2 are excluded from operator analysis."],
    }
    return facts


def build_ai_report_facts_for_calendar_period(result, period_start, period_end, previous_start, previous_end):
    def within(items, start, end):
        return [
            item for item in items
            if start <= str(item.get("date") or "") <= end
        ]

    production_rows = result.get("production_rows", [])
    operator_rows = result.get("operator_rows", [])
    current_rows = within(production_rows, period_start, period_end)
    current_operators = within(operator_rows, period_start, period_end)
    previous_rows = within(production_rows, previous_start, previous_end)
    previous_operators = within(operator_rows, previous_start, previous_end)
    operators = build_operator_summary(current_operators)
    return {
        "period": {
            "start": period_start,
            "end": period_end,
            "rolling_days": None,
            "calendar_month": True,
        },
        "current": summarize_rows(current_rows, current_operators),
        "previous_period": {
            "start": previous_start,
            "end": previous_end,
            **summarize_rows(previous_rows, previous_operators),
        },
        "operators": [
            {key: safe_number(value) if isinstance(value, float) else value for key, value in item.items()}
            for item in operators
        ],
        "business_context": {
            "preferred_backlog_weeks": 2.0,
            "backlog_interpretation": "Backlog below approximately two weeks means production has less new-order repair work available and should prompt sales activity to secure more orders. Growth toward two weeks is positive, not a turnaround concern. A low backlog normally enables faster-than-usual customer turnaround.",
        },
        "data_notes": [
            "The current and previous periods are complete calendar months.",
            "Today is excluded.",
            "Backlog contains new customer orders awaiting their first repair only; re-cook work is never included in backlog.",
            "The preferred backlog is approximately two weeks. A backlog below that level can constrain production because only received customer mats can be repaired, and should prompt the sales team to secure more orders.",
            "Backlog growth that remains below approximately two weeks is movement toward a healthier workload and must not be described as harming turnaround. Low backlog normally means orders can be returned faster than usual.",
            "Re-cook activity is a separate quality measure and is separate from press throughput.",
            "Press throughput is FF1 + FF2 + FF3 total LF.",
            "Trudy Dunlap, Kelly Bainbridge, Lois Horace, Temp1 and Temp2 are excluded from operator analysis.",
        ],
    }


def period_label(payload):
    rows = payload.get("rows", [])
    return f"{rows[0]['date']} to {rows[-1]['date']}" if rows else "No completed production dates"


def clean_filename(value):
    return re.sub(r"[^A-Za-z0-9._-]+", "-", str(value or "")).strip("-")


def build_ai_analysis_docx(facts, analysis):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65); section.bottom_margin = Inches(0.65); section.left_margin = Inches(0.7); section.right_margin = Inches(0.7)
    styles = document.styles
    styles["Normal"].font.name = "Arial"; styles["Normal"].font.size = Pt(10)
    for name, size, color in (("Title", 24, "245CFF"), ("Heading 1", 15, "1F2933"), ("Heading 2", 12, "245CFF")):
        styles[name].font.name = "Arial"; styles[name].font.size = Pt(size); styles[name].font.color.rgb = RGBColor.from_string(color); styles[name].font.bold = True
    title = document.add_paragraph(style="Title"); title.add_run("Numat Production Analysis")
    subtitle = document.add_paragraph(f"{facts['period']['start']} to {facts['period']['end']} · Completed production days only")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    current = facts["current"]
    document.add_heading("1. Overall production summary", level=1)
    table = document.add_table(rows=2, cols=4); table.style = "Light Shading Accent 1"
    metrics = [("Production days", current["production_days"]), ("Average daily value", f"${current['average_daily_revenue']:,.0f}" if current["average_daily_revenue"] is not None else "—"), ("Average press LF", f"{current['average_press_lf']:,.0f}" if current["average_press_lf"] is not None else "—"), ("Plant productivity", f"{current['plant_productivity']:.1f}%" if current["plant_productivity"] is not None else "—")]
    for col, (label, value) in enumerate(metrics):
        table.cell(0, col).text = label; table.cell(1, col).text = str(value)
        table.cell(0, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; table.cell(1, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    add_section(document, "Executive interpretation", analysis.get("executive_summary", []))
    add_section(document, "2. Previous-period context", analysis.get("comparison", []))
    add_section(document, "3. Press and production flow", analysis.get("press_flow", []))
    add_section(document, "4. Quality and re-cook activity", analysis.get("quality", []))
    add_section(document, "5. Key strengths", analysis.get("strengths", []))
    document.add_page_break()
    add_section(document, "6. Key focus areas", analysis.get("focus_areas", []), numbered=True)
    document.add_heading("Operator performance", level=1)
    op_table = document.add_table(rows=1, cols=6); op_table.style = "Light Shading Accent 1"
    op_table.autofit = False
    column_widths_dxa = [900, 1800, 800, 1500, 1200, 4024]
    for cell, label in zip(op_table.rows[0].cells, ["Rank", "Name", "Days", "Productivity", "LF target", "Comment"]): cell.text = label
    header_properties = op_table.rows[0]._tr.get_or_add_trPr()
    header_properties.append(OxmlElement("w:tblHeader"))
    comments = analysis.get("operator_comments", {}) if isinstance(analysis.get("operator_comments"), dict) else {}
    for rank, item in enumerate(sorted(facts.get("operators", []), key=lambda value: -(value.get("target_achieved") or 0)), start=1):
        cells = op_table.add_row().cells
        values = [rank, item.get("name"), item.get("days"), f"{item.get('productivity'):.1f}%" if item.get("productivity") is not None else "—", f"{item.get('target_achieved'):.1f}%" if item.get("target_achieved") is not None else "—", comments.get(item.get("name"), "")]
        for cell, value in zip(cells, values): cell.text = str(value)
    for row in op_table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
    table_properties = op_table._tbl.tblPr
    table_width = OxmlElement("w:tblW"); table_width.set(qn("w:type"), "dxa"); table_width.set(qn("w:w"), str(sum(column_widths_dxa))); table_properties.append(table_width)
    table_layout = OxmlElement("w:tblLayout"); table_layout.set(qn("w:type"), "fixed"); table_properties.append(table_layout)
    grid = op_table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in column_widths_dxa:
        grid_column = OxmlElement("w:gridCol"); grid_column.set(qn("w:w"), str(width)); grid.append(grid_column)
    for row in op_table.rows:
        for cell, width in zip(row.cells, column_widths_dxa):
            cell_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            cell_width.set(qn("w:type"), "dxa"); cell_width.set(qn("w:w"), str(width))
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
    add_section(document, "Workforce observations", analysis.get("workforce_observations", []))
    document.add_heading("Final message", level=1)
    document.add_paragraph(str(analysis.get("final_message") or "The verified production figures are presented above for management review."))
    footer = section.footer.paragraphs[0]
    footer.text = "AI interpretation uses verified Production Analysis figures only; review recommendations alongside operational context."
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].italic = True
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(91, 107, 124)
    output = BytesIO(); document.save(output); return output.getvalue()


def add_section(document, heading, content, numbered=False):
    document.add_heading(heading, level=1)
    items = content if isinstance(content, list) else [content]
    for item in items:
        text = str(item or "").strip()
        if text:
            document.add_paragraph(text, style="List Number" if numbered else "List Bullet")
